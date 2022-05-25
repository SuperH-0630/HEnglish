import configure
import os

env_dict = os.environ
henglish_conf = env_dict.get("HENGLISH_CONF")
if henglish_conf is not None:
    configure.configure(henglish_conf, encoding="utf-8")


try:
    from core.db import WordDatabase, word
    from docx import Document, shared, enum
    from docx.enum import section, text
    from docx.oxml.ns import qn
except ImportError:
    print("Run `pip install python-docx` first.")
else:
    def make_word(title: str, little_title: str, db: WordDatabase, font0=15, font1=8, font2=7, start=0, end=300,
                  file_name="English"):
        doc = Document()
        sec = doc.sections[0]  # sections对应文档中的“节”

        normal_style = doc.styles['Normal']
        normal_style.font.name = u'微软雅黑'  # 必须先设置font.name
        normal_style.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        normal_style.font.color.rgb = shared.RGBColor(0, 0, 0)

        h1_style = doc.styles['Heading 1']
        h1_style.font.name = u'微软雅黑'  # 必须先设置font.name
        h1_style.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        h1_style.font.size = shared.Pt(30)
        h1_style.font.color.rgb = shared.RGBColor(0, 0, 0)

        h2_style = doc.styles['Heading 2']
        h2_style.font.name = u'微软雅黑'  # 必须先设置font.name
        h2_style.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        h2_style.font.size = shared.Pt(font0)
        h2_style.font.color.rgb = shared.RGBColor(0, 0, 0)

        sec.left_margin = shared.Cm(3.18)  # 以下依次设置左、右、上、下页面边距
        sec.right_margin = shared.Cm(3.18)
        sec.top_margin = shared.Cm(1.27)
        sec.bottom_margin = shared.Cm(1.27)

        sec.orientation = section.WD_ORIENTATION.LANDSCAPE  # 设置纵向
        sec.page_width = shared.Cm(21)  # 设置页面宽度
        sec.page_height = shared.Cm(29.7)  # 设置页面高度
        sec.different_first_page_header_footer = True

        title_par = doc.add_heading(level=1)
        title_par.paragraph_format.alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_par.add_run(title)
        title_run.add_break(text.WD_BREAK.PAGE)

        header = sec.header.paragraphs[0]  # 获取第一个节的页眉
        header.paragraph_format.alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER
        l_title = header.add_run(little_title)  # 添加页面内容
        l_title.font.size = shared.Pt(6)

        footer = sec.footer.paragraphs[0]  # 获取第一个节的页眉
        footer.paragraph_format.alignment = text.WD_PARAGRAPH_ALIGNMENT.CENTER
        power_bt = footer.add_run(f"Power HEnglish! https://github.com/SuperH-0630/HEnglish")  # 添加页面内容
        power_bt.font.size = shared.Pt(6)

        def add_word(w: word.Word):
            doc.add_heading(w.name, level=2)  # 添加标题

            for i in w.comment:
                c: word.Word.Comment = w.comment[i]

                paragraph = doc.add_paragraph()
                part = paragraph.add_run(c.part + "  ")
                part.italic = True
                part.bold = True
                part.font.size = shared.Pt(font1)

                english = paragraph.add_run(c.english)
                english.bold = True
                english.font.size = shared.Pt(font1)

                chinese = paragraph.add_run(c.chinese)
                chinese.bold = True
                chinese.font.size = shared.Pt(font1)

                for eg_index, e in enumerate(c.eg):
                    e = e.split("##")
                    english_sentence = paragraph.add_run(f"Eg {eg_index + 1}: " + e[0])
                    english_sentence.font.size = shared.Pt(font2)

                    chinese_sentence = paragraph.add_run("  " + e[1] + "")
                    chinese_sentence.font.size = shared.Pt(font2)

                ph_format = paragraph.paragraph_format
                ph_format.line_spacing_rule = text.WD_LINE_SPACING.SINGLE  # 设置行间距

        last_ten = -1
        index = start
        while True:
            w = db.find_word_by_index(index)
            if w is None:
                break
            add_word(w)
            index += 1

            if index // 10 > last_ten:
                last_ten = index // 10
                print(index)

            if index == end:
                break

        doc.save(f'{file_name}.docx')
        return index

    word_db = WordDatabase("high_school", configure.conf["DB_TEMPLATE"])

    page = 0
    step = 500
    book = 1
    while True:
        print(f"Book: {book}")
        ret = make_word(f"\n\n\nHigh School English\nWord - {book}",
                        f"High School English Word - {book}",
                        word_db,
                        start=page,
                        end=page + step,
                        file_name=f"English{book}")
        if ret != page + step:
            break
        page += step
        book += 1
